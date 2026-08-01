from __future__ import annotations

import io
import json
import re
from dataclasses import dataclass
from datetime import date, datetime
from html import escape
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import (
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

META_PREFIX = "__DIGIT_META_V1__"

MONTHS_ES = (
    "enero",
    "febrero",
    "marzo",
    "abril",
    "mayo",
    "junio",
    "julio",
    "agosto",
    "septiembre",
    "octubre",
    "noviembre",
    "diciembre",
)

DOCUMENT_LABELS: dict[str, str] = {
    "certificado_trabajo_a": "Certificado de Trabajo",
    "certificado_trabajo_b": "Constancia Laboral",
    "constancia": "Constancia de Trabajo",
    "aguinaldo_anual": "Recibo de Pago de Aguinaldo",
    "aguinaldo_proporcional": "Liquidación de Aguinaldo Proporcional",
    "permiso_paternidad": "Solicitud de Permiso por Paternidad",
    "contrato_trabajo": "Contrato de Trabajo",
    "ficha_empleado": "Ficha de Empleado",
    "solicitud_vacacion": "Solicitud de Vacaciones",
    "usufructo_vacaciones": "Constancia de Usufructo de Vacaciones",
    "notificacion_preaviso": "Notificación de Preaviso",
    "renuncia": "Nota de Renuncia",
    "despido": "Comunicación de Despido",
}

SPECIAL_NUMBERS = {
    0: "cero",
    1: "uno",
    2: "dos",
    3: "tres",
    4: "cuatro",
    5: "cinco",
    6: "seis",
    7: "siete",
    8: "ocho",
    9: "nueve",
    10: "diez",
    11: "once",
    12: "doce",
    13: "trece",
    14: "catorce",
    15: "quince",
    16: "dieciséis",
    17: "diecisiete",
    18: "dieciocho",
    19: "diecinueve",
    20: "veinte",
    21: "veintiuno",
    22: "veintidós",
    23: "veintitrés",
    24: "veinticuatro",
    25: "veinticinco",
    26: "veintiséis",
    27: "veintisiete",
    28: "veintiocho",
    29: "veintinueve",
}
TENS = {30: "treinta", 40: "cuarenta", 50: "cincuenta", 60: "sesenta", 70: "setenta", 80: "ochenta", 90: "noventa"}
HUNDREDS = {
    100: "cien",
    200: "doscientos",
    300: "trescientos",
    400: "cuatrocientos",
    500: "quinientos",
    600: "seiscientos",
    700: "setecientos",
    800: "ochocientos",
    900: "novecientos",
}


@dataclass(frozen=True)
class ExportData:
    document_type: str
    title: str
    city: str
    issue_date: date
    company_name: str
    company_ruc: str = ""
    company_address: str = ""
    company_phone: str = ""
    company_email: str = ""
    legal_representative: str = ""
    employee_name: str = ""
    employee_document: str = ""
    position: str = ""
    admission_date: date | None = None
    salary: int = 0
    body: str = ""
    status: str = "Borrador"
    created_by: str = ""
    created_at: datetime | None = None
    metadata: dict[str, Any] | None = None


def format_date_long_es(value: date | datetime | None) -> str:
    if not value:
        return ""
    return f"{value.day} de {MONTHS_ES[value.month - 1]} de {value.year}"


def format_date_short(value: date | datetime | None) -> str:
    if not value:
        return ""
    return value.strftime("%d/%m/%Y")


def format_gs(value: int | float | None) -> str:
    return f"{int(round(value or 0)):,}".replace(",", ".")


def _apocopate(text: str) -> str:
    if text.endswith("veintiuno"):
        return text[:-9] + "veintiún"
    if text.endswith(" y uno"):
        return text[:-6] + " y un"
    if text.endswith("uno"):
        return text[:-3] + "un"
    return text


def _under_thousand(number: int) -> str:
    if number in SPECIAL_NUMBERS:
        return SPECIAL_NUMBERS[number]
    if number < 100:
        tens = (number // 10) * 10
        unit = number % 10
        return TENS[tens] if unit == 0 else f"{TENS[tens]} y {SPECIAL_NUMBERS[unit]}"
    hundreds = (number // 100) * 100
    remainder = number % 100
    prefix = "ciento" if hundreds == 100 and remainder else HUNDREDS[hundreds]
    return prefix if remainder == 0 else f"{prefix} {_under_thousand(remainder)}"


def number_to_words_es(number: int) -> str:
    number = int(number)
    if number < 0:
        return "menos " + number_to_words_es(abs(number))
    if number < 1000:
        return _under_thousand(number)
    if number < 1_000_000:
        thousands, remainder = divmod(number, 1000)
        prefix = "mil" if thousands == 1 else f"{_apocopate(number_to_words_es(thousands))} mil"
        return prefix if remainder == 0 else f"{prefix} {number_to_words_es(remainder)}"
    if number < 1_000_000_000:
        millions, remainder = divmod(number, 1_000_000)
        prefix = "un millón" if millions == 1 else f"{_apocopate(number_to_words_es(millions))} millones"
        return prefix if remainder == 0 else f"{prefix} {number_to_words_es(remainder)}"
    billions, remainder = divmod(number, 1_000_000_000)
    prefix = "mil millones" if billions == 1 else f"{_apocopate(number_to_words_es(billions))} mil millones"
    return prefix if remainder == 0 else f"{prefix} {number_to_words_es(remainder)}"


def guaranies_in_words(value: int) -> str:
    value = max(0, int(value or 0))
    return f"{_apocopate(number_to_words_es(value))} guaraníes"


def encode_metadata(notes: str = "", **metadata: Any) -> str:
    payload = {"notes": notes.strip(), **metadata}
    return META_PREFIX + json.dumps(payload, ensure_ascii=False, separators=(",", ":"))


def decode_metadata(raw: str | None) -> dict[str, Any]:
    if not raw:
        return {"notes": ""}
    if not raw.startswith(META_PREFIX):
        return {"notes": raw}
    try:
        payload = json.loads(raw[len(META_PREFIX) :])
        return payload if isinstance(payload, dict) else {"notes": ""}
    except (json.JSONDecodeError, TypeError):
        return {"notes": raw}


def optional_date(value: Any) -> date | None:
    if isinstance(value, date):
        return value
    if isinstance(value, str) and value:
        try:
            return date.fromisoformat(value)
        except ValueError:
            return None
    return None


def build_document_body(
    document_type: str,
    *,
    company_name: str,
    employee_name: str,
    employee_document: str,
    position: str,
    admission_date: date | None,
    salary: int,
    issue_date: date,
    city: str,
    metadata: dict[str, Any],
) -> tuple[str, str]:
    if document_type not in DOCUMENT_LABELS:
        raise ValueError("Tipo de documento no soportado")

    title = DOCUMENT_LABELS[document_type]
    notes = str(metadata.get("notes", "")).strip()
    identity = employee_name
    if employee_document:
        identity += f", con C.I. N.º {employee_document}"
    role = position or "el cargo registrado en su legajo"
    admission = format_date_long_es(admission_date) or "la fecha registrada en su legajo"
    salary_words = guaranies_in_words(salary)
    salary_display = f"Gs. {format_gs(salary)} ({salary_words})" if salary else "la remuneración registrada en su legajo"
    period_start = optional_date(metadata.get("period_start"))
    period_end = optional_date(metadata.get("period_end"))
    effective_date = optional_date(metadata.get("effective_date"))
    leave_start = optional_date(metadata.get("leave_start"))
    leave_end = optional_date(metadata.get("leave_end"))
    amount = int(metadata.get("amount") or salary or 0)
    recipient = str(metadata.get("recipient") or "Encargado/a de Recursos Humanos").strip()
    nationality = str(metadata.get("nationality") or "paraguaya").strip()
    civil_status = str(metadata.get("civil_status") or "").strip()

    bodies: dict[str, str] = {
        "certificado_trabajo_a": (
            f"Por medio de la presente, {company_name} certifica que {identity} presta servicios en la empresa "
            f"en el cargo de {role}, desde el {admission}, percibiendo actualmente una remuneración mensual de {salary_display}.\n\n"
            "Se expide el presente certificado a solicitud de la persona interesada, para los fines que estime convenientes."
        ),
        "certificado_trabajo_b": (
            f"Se deja constancia de que {identity} forma parte del plantel de {company_name}, desempeñándose como {role} "
            f"desde el {admission}. La remuneración mensual registrada es de {salary_display}.\n\n"
            "La presente constancia se emite a pedido de la persona interesada."
        ),
        "constancia": (
            f"{company_name} hace constar que {identity}, de nacionalidad {nationality}"
            + (f", de estado civil {civil_status}" if civil_status else "")
            + f", presta servicios en la empresa en el cargo de {role}, desde el {admission}, y percibe una remuneración mensual de {salary_display}.\n\n"
            f"Se expide la presente constancia en {city}, a los {format_date_long_es(issue_date)}, para los fines que correspondan."
        ),
        "aguinaldo_anual": (
            f"El empleador {company_name} deja constancia de haber abonado a {identity} la suma de Gs. {format_gs(amount)} "
            f"({guaranies_in_words(amount)}), en concepto de aguinaldo correspondiente al período comprendido entre el "
            f"{format_date_short(period_start)} y el {format_date_short(period_end)}, calculado conforme al artículo 243 del Código del Trabajo.\n\n"
            "Con la firma del presente documento, el trabajador declara haber recibido el importe indicado."
        ),
        "aguinaldo_proporcional": (
            f"El empleador {company_name} deja constancia de haber abonado a {identity} la suma de Gs. {format_gs(amount)} "
            f"({guaranies_in_words(amount)}), en concepto de aguinaldo proporcional devengado desde el "
            f"{format_date_short(period_start)} hasta el {format_date_short(period_end)}, conforme al artículo 244 del Código del Trabajo.\n\n"
            "Con la firma del presente documento, el trabajador declara haber recibido el importe indicado."
        ),
        "permiso_paternidad": (
            f"Señor/a\n{recipient}\n{company_name}\n\nRef.: Solicitud de permiso por paternidad\n\n"
            f"Yo, {identity}, trabajador de la empresa, solicito el permiso por paternidad correspondiente a dos semanas posteriores al parto, "
            f"con goce de sueldo, desde el {format_date_short(leave_start)} hasta el {format_date_short(leave_end)}, de conformidad con el artículo 13, inciso b), de la Ley N.º 5508/2015.\n\n"
            "Me comprometo a presentar la documentación respaldatoria requerida para el registro del permiso. Solicito se deje constancia de la recepción de esta nota."
        ),
        "contrato_trabajo": (
            f"BORRADOR PARA REVISIÓN PROFESIONAL.\n\nEntre {company_name}, en carácter de empleador, y {identity}, en carácter de trabajador/a, "
            f"se prepara el presente borrador de contrato para el cargo de {role}, con inicio el {admission} y remuneración mensual de {salary_display}.\n\n"
            "Las condiciones de jornada, funciones, lugar de trabajo, descansos, beneficios, duración y terminación deberán completarse y revisarse antes de la firma."
        ),
        "ficha_empleado": (
            f"EMPRESA: {company_name}\nFUNCIONARIO/A: {employee_name}\nCÉDULA: {employee_document or '—'}\nCARGO: {role}\n"
            f"FECHA DE INGRESO: {admission}\nSALARIO REGISTRADO: {salary_display}"
        ),
        "solicitud_vacacion": (
            f"Yo, {identity}, solicito a {company_name} el usufructo de mis vacaciones desde el {format_date_short(leave_start)} "
            f"hasta el {format_date_short(leave_end)}. Las fechas quedarán sujetas a aprobación y constancia escrita de la empresa."
        ),
        "usufructo_vacaciones": (
            f"{company_name} deja constancia de que {identity}, quien se desempeña como {role}, usufructará o ha usufructuado "
            f"sus vacaciones desde el {format_date_short(leave_start)} hasta el {format_date_short(leave_end)}."
        ),
        "notificacion_preaviso": (
            f"Señor/a\n{employee_name}\nC.I. N.º {employee_document}\nPresente\n\n"
            f"Por medio de la presente, {company_name} le comunica el preaviso de terminación de la relación laboral, con fecha efectiva "
            f"{format_date_long_es(effective_date)}. La notificación se cursa por escrito conforme a los artículos 87 y 88 del Código del Trabajo.\n\n"
            "Durante el período de preaviso, y sin disminución de salario, podrá optar por una licencia diaria de dos horas dentro de la jornada legal, "
            "por un día a la semana o por el uso continuado del tiempo correspondiente, para buscar un nuevo empleo, conforme al artículo 89 del Código del Trabajo.\n\n"
            "Se solicita firmar la recepción de la presente, sin que ello implique conformidad con su contenido."
        ),
        "renuncia": (
            f"Señor/a\n{recipient}\n{company_name}\n\nRef.: Comunicación de renuncia\n\n"
            f"Yo, {identity}, comunico mi decisión de dar por terminada la relación laboral con fecha efectiva "
            f"{format_date_long_es(effective_date)}. Solicito se practique la liquidación final y se expida la constancia de trabajo correspondiente."
        ),
        "despido": (
            f"BORRADOR PARA REVISIÓN PROFESIONAL.\n\nPor medio de la presente, {company_name} comunica a {identity} la terminación de la relación laboral "
            f"con fecha efectiva {format_date_long_es(effective_date)}. La causa, liquidación, preaviso y documentación respaldatoria deben revisarse y detallarse antes de su entrega."
        ),
    }
    body = bodies[document_type]
    if notes:
        body += f"\n\nObservaciones: {notes}"
    return title, body


def export_data_from_certificate(certificate: Any) -> ExportData:
    company = certificate.company
    metadata = decode_metadata(certificate.observations)
    return ExportData(
        document_type=certificate.document_type,
        title=certificate.title,
        city=certificate.city,
        issue_date=certificate.issue_date,
        company_name=certificate.company_name_snapshot,
        company_ruc=getattr(company, "ruc", "") or "",
        company_address=getattr(company, "address", "") or "",
        company_phone=getattr(company, "phone", "") or "",
        company_email=getattr(company, "email", "") or "",
        legal_representative=getattr(company, "legal_representative", "") or getattr(company, "responsible_name", "") or "",
        employee_name=certificate.employee_name_snapshot,
        employee_document=certificate.employee_document_snapshot,
        position=certificate.position_snapshot,
        admission_date=certificate.admission_date_snapshot,
        salary=certificate.salary_snapshot,
        body=certificate.body,
        status=certificate.status,
        created_by=certificate.created_by,
        created_at=certificate.created_at,
        metadata=metadata,
    )


def safe_download_name(title: str, employee_name: str, suffix: str) -> str:
    raw = f"{title} - {employee_name}".strip(" -")
    cleaned = re.sub(r"[^A-Za-z0-9ÁÉÍÓÚáéíóúÑñÜü ._-]+", "", raw).strip().replace(" ", "_")
    return f"{cleaned[:120] or 'documento_laboral'}.{suffix}"


def _set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_margins(cell, **kwargs: int) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin in ["top", "start", "bottom", "end"]:
        if margin in kwargs:
            node = tc_mar.find(qn(f"w:{margin}"))
            if node is None:
                node = OxmlElement(f"w:{margin}")
                tc_mar.append(node)
            node.set(qn("w:w"), str(kwargs[margin]))
            node.set(qn("w:type"), "dxa")


def _docx_header(document: Document, data: ExportData) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(2.1)
    table.columns[1].width = Cm(14.7)
    left, right = table.rows[0].cells
    left.width = Cm(2.1)
    right.width = Cm(14.7)
    left.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    _set_cell_shading(left, "173B86")
    _set_cell_margins(left, top=160, start=100, bottom=160, end=100)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("DL")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(255, 255, 255)

    p = right.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(data.company_name.upper())
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(23, 59, 134)
    details = []
    if data.company_ruc:
        details.append(f"RUC {data.company_ruc}")
    if data.company_address:
        details.append(data.company_address)
    if data.company_phone:
        details.append(f"Tel. {data.company_phone}")
    if data.company_email:
        details.append(data.company_email)
    if details:
        p2 = right.add_paragraph(" · ".join(details))
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        for run in p2.runs:
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(85, 97, 116)

    border = document.add_paragraph()
    border.paragraph_format.space_before = Pt(2)
    border.paragraph_format.space_after = Pt(12)
    run = border.add_run("━" * 82)
    run.font.color.rgb = RGBColor(23, 59, 134)
    run.font.size = Pt(6)


def _add_signature_table_docx(document: Document, data: ExportData) -> None:
    document.add_paragraph().paragraph_format.space_after = Pt(18)
    if data.document_type in {"notificacion_preaviso", "aguinaldo_anual", "aguinaldo_proporcional"}:
        table = document.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for idx, label in enumerate(("Firma del trabajador/a", "Firma del empleador/a")):
            cell = table.cell(0, idx)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("\n\n_______________________________\n").bold = False
            r = p.add_run(label)
            r.bold = True
            r.font.size = Pt(9)
            if idx == 0 and data.employee_name:
                rr = p.add_run(f"\n{data.employee_name}")
                rr.font.size = Pt(8.5)
                if data.employee_document:
                    p.add_run(f"\nC.I. N.º {data.employee_document}").font.size = Pt(8.5)
            if idx == 1 and data.legal_representative:
                p.add_run(f"\n{data.legal_representative}").font.size = Pt(8.5)
        if data.document_type == "notificacion_preaviso":
            p = document.add_paragraph("Fecha de recepción: ____/____/________")
            p.paragraph_format.space_before = Pt(14)
            p.runs[0].font.size = Pt(9)
    elif data.document_type in {"permiso_paternidad", "renuncia", "solicitud_vacacion"}:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("\n\n_______________________________\n")
        r = p.add_run(data.employee_name or "Firma del trabajador/a")
        r.bold = True
        if data.employee_document:
            p.add_run(f"\nC.I. N.º {data.employee_document}")
        p2 = document.add_paragraph("RECIBIDO: ____________________    FECHA: ____/____/________")
        p2.paragraph_format.space_before = Pt(14)
        p2.runs[0].font.size = Pt(9)
    else:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(28)
        p.add_run("_______________________________\n")
        r = p.add_run(data.legal_representative or "Firma autorizada")
        r.bold = True
        p.add_run(f"\n{data.company_name}")


def build_docx_bytes(data: ExportData) -> bytes:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(11)
    styles["Normal"].paragraph_format.line_spacing = 1.25

    _docx_header(document, data)

    date_p = document.add_paragraph(f"{data.city}, {format_date_long_es(data.issue_date)}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.paragraph_format.space_after = Pt(18)
    date_p.runs[0].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(18)
    run = title.add_run(data.title.upper())
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = RGBColor(11, 31, 72)

    meta = data.metadata or {}
    document_number = str(meta.get("document_number") or "").strip()
    if document_number:
        number_p = document.add_paragraph(f"Documento N.º {document_number}")
        number_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        number_p.paragraph_format.space_after = Pt(10)
        number_p.runs[0].font.size = Pt(9)
        number_p.runs[0].italic = True

    for block in data.body.split("\n\n"):
        lines = block.split("\n")
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if len(lines) == 1 else WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(10)
        for index, line in enumerate(lines):
            if index:
                p.add_run().add_break()
            run = p.add_run(line)
            run.font.size = Pt(11)
            if line.lower().startswith("ref.:") or line.lower().startswith("ref:"):
                run.bold = True
                run.italic = True

    _add_signature_table_docx(document, data)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Generado por Digit Laboral · {format_date_short(data.created_at or datetime.now())} · Estado: {data.status}"
    )
    fr.font.size = Pt(7.5)
    fr.font.color.rgb = RGBColor(120, 130, 145)

    if data.status.lower() == "borrador":
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run("BORRADOR - REVISAR ANTES DE FIRMAR")
        r.bold = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(170, 70, 70)

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _pdf_footer(canvas, doc, data: ExportData) -> None:
    canvas.saveState()
    width, _ = A4
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#7A8597"))
    canvas.drawCentredString(
        width / 2,
        10 * mm,
        f"Generado por Digit Laboral · {format_date_short(data.created_at or datetime.now())} · Estado: {data.status} · Página {doc.page}",
    )
    canvas.restoreState()


def build_pdf_bytes(data: ExportData) -> bytes:
    stream = io.BytesIO()
    doc = SimpleDocTemplate(
        stream,
        pagesize=A4,
        rightMargin=20 * mm,
        leftMargin=20 * mm,
        topMargin=16 * mm,
        bottomMargin=18 * mm,
        title=data.title,
        author="Digit Laboral",
    )
    styles = getSampleStyleSheet()
    company_style = ParagraphStyle(
        "Company",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        leading=17,
        textColor=colors.HexColor("#173B86"),
        spaceAfter=2,
    )
    detail_style = ParagraphStyle(
        "Detail",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8.3,
        leading=10,
        textColor=colors.HexColor("#556174"),
    )
    date_style = ParagraphStyle("Date", parent=styles["Normal"], fontName="Helvetica", fontSize=10, alignment=TA_RIGHT, spaceAfter=14)
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        leading=17,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#0B1F48"),
        spaceAfter=16,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=10.5,
        leading=16,
        alignment=TA_JUSTIFY,
        spaceAfter=9,
    )
    left_style = ParagraphStyle("Left", parent=body_style, alignment=TA_LEFT)
    small_style = ParagraphStyle("Small", parent=styles["Normal"], fontName="Helvetica", fontSize=8.5, leading=11)
    draft_style = ParagraphStyle(
        "Draft",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=8,
        textColor=colors.HexColor("#A74646"),
        alignment=TA_CENTER,
        spaceBefore=12,
    )

    story = []
    details = []
    if data.company_ruc:
        details.append(f"RUC {escape(data.company_ruc)}")
    if data.company_address:
        details.append(escape(data.company_address))
    if data.company_phone:
        details.append(f"Tel. {escape(data.company_phone)}")
    if data.company_email:
        details.append(escape(data.company_email))
    header = Table(
        [
            [
                Paragraph("<b>DL</b>", ParagraphStyle("DL", parent=styles["Normal"], fontName="Helvetica-Bold", fontSize=18, textColor=colors.white, alignment=TA_CENTER)),
                [Paragraph(escape(data.company_name.upper()), company_style), Paragraph(" · ".join(details), detail_style) if details else ""],
            ]
        ],
        colWidths=[18 * mm, 152 * mm],
    )
    header.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#173B86")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ("LINEBELOW", (0, 0), (-1, -1), 1.3, colors.HexColor("#173B86")),
            ]
        )
    )
    story.extend([header, Spacer(1, 9 * mm)])
    story.append(Paragraph(f"{escape(data.city)}, {escape(format_date_long_es(data.issue_date))}", date_style))
    story.append(Paragraph(escape(data.title.upper()), title_style))
    document_number = str((data.metadata or {}).get("document_number") or "").strip()
    if document_number:
        story.append(Paragraph(f"Documento N.º {escape(document_number)}", ParagraphStyle("Number", parent=small_style, alignment=TA_RIGHT, spaceAfter=8)))

    for block in data.body.split("\n\n"):
        lines = block.split("\n")
        content = "<br/>".join(escape(line) for line in lines)
        story.append(Paragraph(content, left_style if len(lines) > 1 else body_style))

    story.append(Spacer(1, 17 * mm))
    if data.document_type in {"notificacion_preaviso", "aguinaldo_anual", "aguinaldo_proporcional"}:
        left = "______________________________<br/><b>Firma del trabajador/a</b>"
        if data.employee_name:
            left += f"<br/>{escape(data.employee_name)}"
        if data.employee_document:
            left += f"<br/>C.I. N.º {escape(data.employee_document)}"
        right = "______________________________<br/><b>Firma del empleador/a</b>"
        if data.legal_representative:
            right += f"<br/>{escape(data.legal_representative)}"
        signatures = Table(
            [[Paragraph(left, ParagraphStyle("SigL", parent=small_style, alignment=TA_CENTER)), Paragraph(right, ParagraphStyle("SigR", parent=small_style, alignment=TA_CENTER))]],
            colWidths=[83 * mm, 83 * mm],
        )
        signatures.setStyle(TableStyle([("VALIGN", (0, 0), (-1, -1), "TOP")]))
        story.append(signatures)
        if data.document_type == "notificacion_preaviso":
            story.append(Spacer(1, 9 * mm))
            story.append(Paragraph("Fecha de recepción: ____/____/________", small_style))
    elif data.document_type in {"permiso_paternidad", "renuncia", "solicitud_vacacion"}:
        sig = "______________________________<br/>"
        sig += f"<b>{escape(data.employee_name or 'Firma del trabajador/a')}</b>"
        if data.employee_document:
            sig += f"<br/>C.I. N.º {escape(data.employee_document)}"
        story.append(Paragraph(sig, ParagraphStyle("EmployeeSig", parent=small_style, alignment=TA_CENTER)))
        story.append(Spacer(1, 9 * mm))
        story.append(Paragraph("RECIBIDO: ____________________    FECHA: ____/____/________", small_style))
    else:
        sig = "______________________________<br/>"
        sig += f"<b>{escape(data.legal_representative or 'Firma autorizada')}</b><br/>{escape(data.company_name)}"
        story.append(Paragraph(sig, ParagraphStyle("EmployerSig", parent=small_style, alignment=TA_CENTER)))

    if data.status.lower() == "borrador":
        story.append(Paragraph("BORRADOR - REVISAR ANTES DE FIRMAR", draft_style))

    doc.build(story, onFirstPage=lambda c, d: _pdf_footer(c, d, data), onLaterPages=lambda c, d: _pdf_footer(c, d, data))
    return stream.getvalue()
